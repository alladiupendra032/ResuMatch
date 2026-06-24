"""
Resume Parsing Engine
Extracts structured data from PDF/DOCX resume files using:
- PyPDF2 / python-docx for text extraction
- Regex for contact info extraction
- Keyword matching for skills, education, certifications
"""

import re
import io
from typing import List, Tuple, Optional
import PyPDF2
import docx

# ─── Skill Keywords Database ────────────────────────────────────────────────────

TECH_SKILLS = {
    # Programming Languages
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "go", "golang",
    "rust", "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "perl",
    "shell", "bash", "powershell",

    # Web Frontend
    "react", "react.js", "reactjs", "angular", "vue", "vue.js", "vuejs", "next.js",
    "nextjs", "nuxt", "svelte", "html", "css", "sass", "scss", "tailwind",
    "bootstrap", "jquery", "redux", "graphql", "webpack", "vite",

    # Web Backend
    "node.js", "nodejs", "express", "fastapi", "django", "flask", "spring",
    "spring boot", "laravel", "rails", "asp.net", "nestjs", "hapi",

    # Databases
    "mongodb", "mysql", "postgresql", "postgres", "sqlite", "redis", "elasticsearch",
    "cassandra", "oracle", "sql server", "dynamodb", "firebase", "supabase",

    # Cloud & DevOps
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
    "terraform", "ansible", "jenkins", "github actions", "ci/cd", "linux",
    "nginx", "apache",

    # Data & ML
    "machine learning", "deep learning", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "matplotlib", "seaborn", "tableau",
    "power bi", "data analysis", "nlp", "computer vision", "opencv",

    # Mobile
    "android", "ios", "react native", "flutter", "xamarin",

    # Tools & Others
    "git", "github", "gitlab", "jira", "agile", "scrum", "rest api", "restful",
    "microservices", "graphql", "kafka", "rabbitmq", "celery", "oauth",
    "jwt", "websocket", "grpc", "postman", "selenium", "pytest", "jest",
    "webpack", "babel", "npm", "yarn", "pip", "maven", "gradle",
}

# Common certification keywords
CERTIFICATION_KEYWORDS = [
    "aws certified", "azure certified", "google certified", "gcp certified",
    "pmp", "prince2", "scrum master", "csm", "safe", "cissp", "ceh",
    "comptia", "cisco", "ccna", "ccnp", "oracle certified", "java certified",
    "microsoft certified", "tensorflow certified", "google analytics",
    "hubspot", "salesforce", "data science", "machine learning certificate",
]

# Education degree keywords mapped to levels
EDUCATION_LEVELS = {
    "high school": 1, "secondary": 1, "ssc": 1, "hsc": 1,
    "associate": 2, "diploma": 2,
    "bachelor": 3, "b.tech": 3, "b.sc": 3, "b.e": 3, "bsc": 3, "be": 3,
    "b.com": 3, "bca": 3, "bba": 3, "undergraduate": 3,
    "master": 4, "m.tech": 4, "m.sc": 4, "m.e": 4, "msc": 4, "mba": 4,
    "mca": 4, "m.com": 4, "postgraduate": 4, "pg diploma": 3,
    "phd": 5, "ph.d": 5, "doctorate": 5, "doctoral": 5,
}

# Section headers used to split resume into sections
SECTION_HEADERS = {
    "experience": ["experience", "work experience", "employment", "work history", "professional experience"],
    "education": ["education", "academic background", "qualifications", "academic qualifications"],
    "skills": ["skills", "technical skills", "core competencies", "technologies", "expertise"],
    "certifications": ["certifications", "certificates", "licenses", "accreditations"],
    "projects": ["projects", "personal projects", "academic projects", "notable projects"],
    "summary": ["summary", "objective", "profile", "about me", "professional summary"],
}


# ─── Text Extraction ────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract raw text from a PDF file."""
    text = ""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract raw text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return text


# ─── Contact Info Extraction ────────────────────────────────────────────────────

def extract_email(text: str) -> str:
    pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    matches = re.findall(pattern, text)
    return matches[0] if matches else ""


def extract_phone(text: str) -> str:
    pattern = r"(?:\+?\d{1,3}[\-\.\s]?)?(?:\(?\d{3}\)?[\-\.\s]?)?\d{3}[\-\.\s]?\d{4}"
    matches = re.findall(pattern, text)
    # Filter out short/invalid matches
    valid = [m.strip() for m in matches if len(re.sub(r"\D", "", m)) >= 10]
    return valid[0] if valid else ""


def extract_name(text: str) -> str:
    """Heuristically extract name from the first non-empty lines of the resume."""
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    for line in lines[:5]:
        # Likely a name if: 2-4 words, no digits, no special chars except hyphen
        words = line.split()
        if 2 <= len(words) <= 4 and all(re.match(r"^[A-Za-z\-]+$", w) for w in words):
            return line
    return lines[0] if lines else ""


# ─── Skills Extraction ─────────────────────────────────────────────────────────

def extract_skills(text: str) -> List[str]:
    """Match tech skills using a keyword list against the resume text."""
    text_lower = text.lower()
    found_skills = set()

    for skill in TECH_SKILLS:
        # Use word-boundary matching
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text_lower):
            found_skills.add(skill.title() if len(skill) > 3 else skill.upper())

    return sorted(list(found_skills))


# ─── Education Extraction ───────────────────────────────────────────────────────

def extract_education(text: str) -> Tuple[List[dict], str, int]:
    """
    Extract education entries and determine the highest education level.
    Returns (education_list, highest_degree_text, level_ordinal).
    """
    education_entries = []
    highest_level = 0
    highest_degree = ""

    lines = text.split("\n")
    for line in lines:
        line_lower = line.lower()
        for degree, level in EDUCATION_LEVELS.items():
            if degree in line_lower:
                entry = {
                    "degree": line.strip(),
                    "field_of_study": "",
                    "institution": "",
                    "year": None,
                }
                # Try to extract year
                year_match = re.search(r"\b(19|20)\d{2}\b", line)
                if year_match:
                    entry["year"] = int(year_match.group())

                education_entries.append(entry)
                if level > highest_level:
                    highest_level = level
                    highest_degree = degree
                break

    return education_entries, highest_degree, highest_level


# ─── Experience Extraction ─────────────────────────────────────────────────────

def extract_experience_years(text: str) -> float:
    """
    Estimate total years of experience from the resume text.
    Looks for patterns like 'X years of experience', or counts date ranges.
    """
    # Direct mention pattern
    direct_patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*years?\s*of\s+(?:professional\s+)?experience",
        r"(\d+(?:\.\d+)?)\+?\s*years?\s+experience",
        r"experience[:\s]+(\d+(?:\.\d+)?)\+?\s*years?",
    ]
    for pattern in direct_patterns:
        match = re.search(pattern, text.lower())
        if match:
            return float(match.group(1))

    # Count year ranges (e.g. 2019–2022, 2020-present)
    year_range_pattern = r"\b(20\d{2}|19\d{2})\s*[-–—to]+\s*(20\d{2}|19\d{2}|present|current|now)\b"
    matches = re.findall(year_range_pattern, text.lower())
    total_years = 0.0
    import datetime
    current_year = datetime.datetime.now().year
    for start, end in matches:
        try:
            start_yr = int(start)
            end_yr = current_year if end in ("present", "current", "now") else int(end)
            if end_yr > start_yr:
                total_years += end_yr - start_yr
        except ValueError:
            pass

    return min(round(total_years, 1), 40.0)  # cap at 40


# ─── Certifications Extraction ─────────────────────────────────────────────────

def extract_certifications(text: str) -> List[str]:
    """Extract certification mentions from the resume."""
    text_lower = text.lower()
    found_certs = []
    for cert in CERTIFICATION_KEYWORDS:
        if cert in text_lower:
            found_certs.append(cert.title())
    return list(set(found_certs))


# ─── Projects Extraction ───────────────────────────────────────────────────────

def extract_projects(text: str) -> List[dict]:
    """Extract project entries from Projects section."""
    projects = []
    lines = text.split("\n")
    in_projects_section = False

    for line in lines:
        stripped = line.strip()
        stripped_lower = stripped.lower()

        # Detect section start
        if any(h in stripped_lower for h in SECTION_HEADERS["projects"]):
            in_projects_section = True
            continue

        # Detect section end (another section header)
        if in_projects_section and stripped_lower and any(
            any(h in stripped_lower for h in headers)
            for section, headers in SECTION_HEADERS.items()
            if section != "projects"
        ):
            in_projects_section = False

        if in_projects_section and stripped and len(stripped) > 5:
            # Basic heuristic: treat each non-empty line as a project entry
            if not projects or len(projects[-1].get("description", "")) > 0:
                projects.append({"title": stripped, "description": "", "technologies": []})
            else:
                projects[-1]["description"] = stripped

    return projects[:5]  # Return max 5 projects


# ─── Main Parsing Pipeline ──────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Main resume parsing pipeline.
    Extracts structured information from a PDF or DOCX resume.
    Returns a dict matching the CandidateProfile schema.
    """
    # 1. Extract text
    if filename.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename.lower().endswith((".docx", ".doc")):
        text = extract_text_from_docx(file_bytes)
    else:
        text = file_bytes.decode("utf-8", errors="ignore")

    if not text.strip():
        return {}

    # 2. Extract fields
    name = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)
    skills = extract_skills(text)
    education_entries, highest_degree, education_level = extract_education(text)
    experience_years = extract_experience_years(text)
    certifications = extract_certifications(text)
    projects = extract_projects(text)

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": skills,
        "experience_years": experience_years,
        "education": education_entries,
        "experience_details": [],
        "certifications": certifications,
        "projects": projects,
        "education_level": education_level,
        "highest_degree": highest_degree,
    }
